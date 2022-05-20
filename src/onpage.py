"""
On page optimization brief.

This module collects metadata from a page and the SERPs for a set of search terms.
It then constructs a content brief in .docx format contaning the the metadata.

References:

> https://python-docx.readthedocs.io/en/latest/
> https://webkul.com/blog/create-word-document-in-python-odoo-python-docx/
> https://mlhive.com/2022/03/create-and-modify-word-docx-files-using-python-docx
"""

import datetime
import sys
from pathlib import Path
from statistics import mean
from typing import Any, Dict, List, Union

import lxml.html
import requests
from docx import Document

from src.apicalls import query_mobile_friendliness, query_pagespeed, query_valueserp
from src.constants import USER_AGENT
from src.formatting import flatten_list, fprint
from src.kwlist import collect_related

session = requests.Session()


def get_response(url: str) -> requests.Response:
    fprint("info", f"fetching {url}")
    try:
        response = session.get(url, headers=USER_AGENT)
    except ConnectionError:
        print("[error] connection error for the url")
    if response and not response.content:
        sys.exit("[error] no data found for the url")
    return response


def get_titles(response: requests.Response) -> List[str]:
    html = lxml.html.fromstring(response.content)
    title_list = html.xpath("//title//text()")
    return [i.strip() for i in title_list]


def get_descriptions(response: requests.Response) -> List[str]:
    html = lxml.html.fromstring(response.content)
    description_list = html.xpath("//meta[@name='description']/@content")
    return [i.strip() for i in description_list]


def get_h1s(response: requests.Response) -> List[str]:
    html = lxml.html.fromstring(response.content)
    h1_list = html.xpath(".//h1//text()")
    return [i.strip() for i in h1_list]


def get_h2s(response: requests.Response) -> List[str]:
    html = lxml.html.fromstring(response.content)
    h2_list = html.xpath("//h2//text()")
    return [i.strip() for i in h2_list]


def get_canonical_link(response: requests.Response) -> Any:
    html = lxml.html.fromstring(response.content)
    return html.xpath("//link[@rel='canonical']/@href")


def get_alternate_links(response: requests.Response) -> Any:
    html = lxml.html.fromstring(response.content)
    return html.xpath("//link[@rel='alternate']/@href")


def get_hreflang(response: requests.Response) -> Any:
    html = lxml.html.fromstring(response.content)
    return html.xpath("//link[@rel='alternate']/@hreflang")


def get_links(response: requests.Response) -> Any:
    html = lxml.html.fromstring(response.content)
    return html.xpath("//@href")


def get_image_links(response: requests.Response) -> Any:
    html = lxml.html.fromstring(response.content)
    return html.xpath("//img/@src")


def get_metarobots(response: requests.Response) -> Any:
    html = lxml.html.fromstring(response.content)
    return html.xpath("//meta[@name='robots']/@content")


def get_paragraphs(response: requests.Response) -> Any:
    html = lxml.html.fromstring(response.content)
    return html.xpath("//p//text()")


def get_wordcount(response: requests.Response) -> int:
    title = " ".join(get_titles(response))
    h1s = " ".join(get_h1s(response))
    h2s = " ".join(get_h2s(response))
    paragraphs = " ".join(get_paragraphs(response))
    all_text = title + h1s + h2s + paragraphs
    return len(all_text.split())


def get_jsonld(response: requests.Response) -> Any:
    html = lxml.html.fromstring(response.content)
    jsonscripts = html.xpath("//script[@type='application/ld+json']//text()")
    if jsonscripts:
        return jsonscripts[0].replace("\t", "")
    else:
        return


def get_page_metadata(url: str) -> Dict[str, Any]:
    fprint("info", f"collecting page metadata for {url}")
    metadata = dict()
    response = get_response(url)
    metadata.update(
        {
            "status_code": response.status_code,
            "titles": get_titles(response),
            "descriptions": get_descriptions(response),
            "h1s": get_h1s(response),
            "h2s": get_h2s(response),
            "canonical": get_canonical_link(response),
            "alternate_links": get_alternate_links(response),
            "hreflang": get_hreflang(response),
            "links": get_links(response),
            "image_links": get_image_links(response),
            "meta_robots": get_metarobots(response),
            "json_ld": get_jsonld(response),
            "wordcount": get_wordcount(response),
        }
    )
    return metadata


def strip_domain(url: str) -> str:
    # strip domain and use it to create the fpath name together with the uuid
    pass


#    for i in pagedata["titles"]:
#        document.add_paragraph(i, style="List Bullet")


def create_brief(url, term, terms, pagedata, serpsdata, fpath):
    document = Document()
    section = document.sections[0]
    header = section.header
    header.paragraphs[0].text = url
    document.add_heading("Content Brief", 0)
    timestamp = datetime.datetime.now()
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run("Date: ").bold = True
    p.add_run(f"{timestamp.strftime('%Y-%m-%d')}")
    p = document.add_paragraph()
    p.add_run("Author: ").bold = True
    document.add_heading("Keyword Information", 1)
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run("Primary Query: ").bold = True
    p.add_run(f"{term}")
    p = document.add_paragraph()
    p.add_run("Secondary Queries: ").bold = True
    p.add_run(f"{', '.join(terms)}")
    p = document.add_paragraph()
    p.add_run("Comments:").bold = True
    document.add_page_break()
    document.add_heading("Recommendations", 1)
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run("Recommended Title:").bold = True
    p = document.add_paragraph()
    p.add_run("Recommended Description:").bold = True
    p = document.add_paragraph()
    p.add_run("Recommended H1s:").bold = True
    p = document.add_paragraph()
    p.add_run("Recommended H2s:").bold = True
    p = document.add_paragraph()
    p.add_run("Recommended Structured Data: ").bold = True
    p = document.add_paragraph()
    p.add_run("Recommended Word Count: ").bold = True
    p.add_run(f"{round(mean(serpsdata['wordcounts']))}")
    document.add_page_break()
    document.add_heading("Technical Health Checks", 1)
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run("Mobile Friendliness: ").bold = True
    p.add_run(f"{query_mobile_friendliness(url)}")
    p = document.add_paragraph()
    p.add_run("Page Speed: ").bold = True
    p.add_run(f"{query_pagespeed(url)}")
    document.add_page_break()
    document.add_heading("Search Intent & Value Assessment", 1)
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run("Competitors:").bold = True
    for i in serpsdata["urls"]:
        document.add_paragraph(i)
    p = document.add_paragraph()
    p.add_run("Search Intent:").bold = True
    p = document.add_paragraph()
    p.add_run("Value To Provide:").bold = True
    document.add_page_break()
    document.add_heading("Page Meta Information", 1)
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run("Current Page Titles:").bold = True
    for i in pagedata["titles"]:
        document.add_paragraph(i)
    p = document.add_paragraph()
    p.add_run("Current Page Descriptions:").bold = True
    for i in pagedata["descriptions"]:
        document.add_paragraph(i)
    p = document.add_paragraph()
    p.add_run("Current H1s:").bold = True
    for i in pagedata["h1s"]:
        document.add_paragraph(i)
    p = document.add_paragraph()
    p.add_run("Current H2s:").bold = True
    for i in pagedata["h2s"]:
        document.add_paragraph(i)
    p = document.add_paragraph()
    document.add_page_break()
    document.add_heading("SERPs Meta Information", 1)
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run("SERPs Titles:").bold = True
    for i in serpsdata["titles"]:
        document.add_paragraph(i)
    p = document.add_paragraph()
    p.add_run("SERPs Descriptions:").bold = True
    for i in serpsdata["descriptions"]:
        document.add_paragraph(i)
    p = document.add_paragraph()
    p.add_run("SERPs H1s:").bold = True
    for i in serpsdata["h1s"]:
        document.add_paragraph(i)
    p = document.add_paragraph()
    p.add_run("SERPs H2s:").bold = True
    for i in serpsdata["h2s"]:
        document.add_paragraph(i)
    p = document.add_paragraph()
    document.save(fpath)


def compile_onpage(url, term, terms, location, fpath):
    fprint("info", f"running onpage analysis for {url} - location: {location}")
    if terms:
        terms.append(term)
        queries = [t.replace(" ", "+") for t in terms]
    else:
        queries = [term]
    organic_results = list()
    for i in queries:
        fprint("info", f"collecting serps for {i}")
        serps = query_valueserp(i, location).json()
        related_results = collect_related(serps)
        for i in serps["organic_results"]:
            if "domain" in i:
                organic_results.append(i["link"])
    page_metadata = get_page_metadata(url)
    page_metadata.update({"related_searches": related_results})
    serps_metadata = dict()
    serps_titles = list()
    serps_descriptions = list()
    serps_h1s = list()
    serps_h2s = list()
    serps_wordcounts = list()
    for i in organic_results:
        metadata = get_page_metadata(i)
        serps_titles.append(metadata["titles"])
        serps_descriptions.append(metadata["descriptions"])
        serps_h1s.append(metadata["h1s"])
        serps_h2s.append(metadata["h2s"])
        serps_wordcounts.append(metadata["wordcount"])
    serps_metadata = {
        "urls": organic_results,
        "titles": flatten_list(serps_titles),
        "descriptions": flatten_list(serps_descriptions),
        "h1s": flatten_list(serps_h1s),
        "h2s": flatten_list(serps_h2s),
        "wordcounts": serps_wordcounts,
    }
    create_brief(url, term, queries, page_metadata, serps_metadata, fpath)
    return {"page_metadata": page_metadata, "serps_metadata": serps_metadata}
